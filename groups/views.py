from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from .models import Group, Lesson, Student, SingleGrade, PersonalNote
from .forms import CreateLessonForm, SingleGradeForm, NameForm, GroupForm, PersonalNoteForm #, TextFieldForm
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from django.urls import reverse
from django.utils import timezone

# Create your views here.

def add_group(request):
    if request.method == 'POST':
        form = GroupForm(request.POST)
        if form.is_valid():
            group = form.save(commit=False)
            group.save()
            return redirect('group_list')
    else:
    	form = GroupForm()
    return render(request, 'groups/add_group.html', {'form':form})

def delete_group(request, lvl):
   group = get_object_or_404(Group, level=lvl)
   if request.method == 'POST':
        group.delete()
        return redirect('group_list')

   return render(request, 'groups/delete_group_confirm.html', {'group': group})
   
def delete_lesson(request, lvl, pk):
    group = get_object_or_404(Group, level=lvl)
    lesson = get_object_or_404(Lesson, pk=pk)
    if request.method == 'POST':
        lesson.delete()
        return redirect('group_lessons', lvl=lvl)
    else:
        form = CreateLessonForm(group=group, instance=lesson)
    return render(request, 'groups/delete_lesson.html', {'lesson': lesson})


def add_student(request, lvl):
   group = get_object_or_404(Group, level=lvl)
   if request.method == 'POST':
       form = NameForm(request.POST)
       if form.is_valid():
           name = form.cleaned_data['name_field']
           surname = form.cleaned_data['surname_field']
                      # Create a new Student instance
           student = Student(name=name, surname=surname, group=group)
           student.save()  # Save the new student to the database
           return redirect('group_lessons', lvl=lvl)
   else:
       form = NameForm()
   return render(request, 'groups/add_student.html', {'form': form, 'group': group})

def delete_student(request, lvl, pk):
    student = get_object_or_404(Student, pk=pk)
    if request.method == 'POST':
        student.delete()
        return redirect('group_lessons', lvl=lvl)

    return render(request, 'groups/delete_student_confirm.html', {'student': student})
"""
def edit_opinion(request, lvl, pk):
    student = get_object_or_404(Student, pk=pk)
    if request.method == 'POST':
        form = TextFieldForm(request.POST)
        if form.is_valid():
            student.opinion = form.cleaned_data['text_field']
            student.save()  # Save the student instance to update the opinion field
            # Process or save the text_value here
            return redirect(reverse('student_info', args=[lvl, pk]))
    else:
        # Pass the current opinion as initial data to the form
        form = TextFieldForm(initial={'text_field': student.opinion})
    return render(request, 'groups/edit_opinion.html', {'form': form, 'student':student})
"""

def edit_opinion(request, lvl, pk):
    student = get_object_or_404(Student, pk=pk)
    group = get_object_or_404(Group, level=lvl)
    if request.method == 'POST':
        form = PersonalNoteForm(request.POST)
        if form.is_valid():
            note = form.save(commit=False)
            note.save()
            return redirect(reverse('student_info', args=[lvl, pk]))
    else:
    	form = PersonalNoteForm()
    	form.fields['student'].initial = student  # Set the initial value for the student field
    	return render(request, 'groups/single_grade.html', {'student':student, 'form':form, 'group':group})

def delete_grade(request, lvl, pk, pk2):
    student = get_object_or_404(Student, pk=pk)
    grade = get_object_or_404(SingleGrade, pk=pk2)

    if request.method == 'POST':
        grade.delete()
        return redirect(reverse('student_info', args=[lvl, pk]))

    return render(request, 'groups/delete_grade_confirm.html', {'student': student, 'grade': grade})

def delete_note(request, lvl, pk, pk2):
    student = get_object_or_404(Student, pk=pk)
    note = get_object_or_404(PersonalNote, pk=pk2)

    if request.method == 'POST':
        note.delete()
        return redirect(reverse('student_info', args=[lvl, pk]))

    return render(request, 'groups/delete_note_confirm.html', {'student': student, 'note': note})    
    
def single_grade(request, lvl, pk):
    student = get_object_or_404(Student, pk=pk)
    group = get_object_or_404(Group, level=lvl)
    if request.method == 'POST':
        form = SingleGradeForm(request.POST)
        if form.is_valid():
            grade = form.save(commit=False)
            grade.save()
            return redirect(reverse('student_info', args=[lvl, pk]))
    else:
    	form = SingleGradeForm()
    	form.fields['student'].initial = student  # Set the initial value for the student field
    	return render(request, 'groups/single_grade.html', {'student':student, 'form':form, 'group':group})

def edit_grade(request, lvl, pk, pk2):
    group = get_object_or_404(Group, level=lvl)
    grade = get_object_or_404(SingleGrade, pk=pk2)
    if request.method == 'POST':
    	form = SingleGradeForm(request.POST, instance=grade)
    	if form.is_valid():
    	    grade = form.save(commit=False)
    	    grade.save()
    	    return redirect(reverse('student_info', args=[lvl, pk]))
    else:
    	form = SingleGradeForm(instance=grade)
    return render(request, 'groups/single_grade.html', {'form':form, 'group':group})


def group_list(request):
    groups = Group.objects.order_by('level')
    return render(request, 'groups/group_list.html', {'groups':groups})

def group_lessons(request, lvl):
    group = get_object_or_404(Group, level=lvl)
    lessons = Lesson.objects.filter(group=group).order_by('-date')
    students_in = Student.objects.filter(group=group).order_by('surname')
    return render(request, 'groups/group_lessons.html', {'group':group, 'lessons':lessons, 'students_in':students_in})

def student_info(request, lvl, pk):
    group = get_object_or_404(Group, level=lvl)
    student = get_object_or_404(Student, pk=pk)
    lessons_absent = student.lessons_absent.all()
    grades = SingleGrade.objects.filter(student=student).order_by('date')
    notes = PersonalNote.objects.filter(student=student).order_by('date')
    return render(request, 'groups/student_info.html', {'student':student, 'lessons_absent':lessons_absent, 'group':group, 'grades':grades, 'notes':notes})

def create_lesson(request, lvl):
    group = get_object_or_404(Group, level=lvl)
    if request.method == 'POST':
        form = CreateLessonForm(request.POST, group=group)
        if form.is_valid():
            lesson = form.save(commit=False)
            lesson.group = group
            lesson.save()
            students_present= form.cleaned_data['students_present']
            all_students = Student.objects.filter(group=group)
            for student in all_students:
                # Check if the student is present in the form's students_present field
                if student in students_present:
                    # If present, add them to the lesson's students_present field
                     lesson.students_present.add(student)
                else:
                    # If absent, add them to the lesson's students_absent field
                    lesson.students_absent.add(student)
            lesson.save()
            return redirect('group_lessons', lvl=lvl)
    else:
        form = CreateLessonForm(group=group)
    return render(request, 'groups/lesson_edit.html', {'form': form, 'group':group})

def edit_lesson(request, lvl, pk):
    group = get_object_or_404(Group, level=lvl)
    lesson = get_object_or_404(Lesson, pk=pk)
    if request.method == 'POST':
        form = CreateLessonForm(request.POST, instance=lesson, group=group)
        if form.is_valid():
            lesson = form.save(commit=False)
            lesson.group = group
            lesson.save()
            students_present= form.cleaned_data['students_present']
            all_students = Student.objects.filter(group=group)
            lesson.students_present.clear()
            lesson.students_absent.clear()
            for student in all_students:
                # Check if the student is present in the form's students_present field
                if student in students_present:
                    # If present, add them to the lesson's students_present field
                     lesson.students_present.add(student)
                else:
                    # If absent, add them to the lesson's students_absent field
                    lesson.students_absent.add(student)
            lesson.save()
            return redirect('group_lessons', lvl=lvl)
    else:
        initial_present_students = lesson.students_present.all()
        form = CreateLessonForm(group=group, instance=lesson)
    return render(request, 'groups/lesson_edit.html', {'form': form, 'group':group})

#___________________________________________________________
#_________________download functions________________________
#___________________________________________________________

def get_status(request, lvl):
    group = get_object_or_404(Group, level=lvl)
    lessons = Lesson.objects.filter(group=group).order_by('date')
    status = Document()
    lesson_table = status.add_table(rows=1, cols=3)
    lesson_table.columns[0].width = Inches(0.3)
    lesson_table.columns[1].width = Inches(1.1)
    lesson_table.columns[2].width = Inches(4.6)
    lesson_table.style = 'Table Grid'
    lesson_table.cell(0, 0).text = '#'
    lesson_table.cell(0, 1).text = 'DATE'
    lesson_table.cell(0, 2).text = 'LESSON PLAN'
    row=0
    for lesson in lessons:
        lesson_table.add_row()
        row +=1
        lesson_table.cell(row, 0).text = str(row)
        lesson_table.cell(row, 1).text = str(lesson.date)
        lesson_table.cell(row, 2).text = str(lesson.subject)
    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = f'inline; filename="Status of {group}.docx"'
    status.save(response)
    return response 

def get_report(request, lvl):
    group = get_object_or_404(Group, level=lvl)
    students = Student.objects.filter(group=group).order_by('surname')
#    lessons = Lesson.objects.filter(group=group).order_by('date')
    report = Document()
    font = report.styles['Normal'].font
    font.name = 'Liberation Serif' 
    default_font_size = Pt(12)
    report.styles['Normal'].font.size = default_font_size
    for student in students:
        p1 = report.add_paragraph(text='Raport postępów w nauce')
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font = p1.runs[0].font
        font.size = Pt(12)
        font.bold = True 
        p2 = report.add_paragraph(text='wrzesień 2023 - 12.11.2023 r.')
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font = p2.runs[0].font
        font.size = Pt(12)
        font.bold = True 
        report.add_paragraph('')
        p3 = report.add_paragraph(text='Imię i nazwisko słuchacza:')
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font = p3.runs[0].font
        font.size = Pt(12)
        report.add_paragraph('')
        p4 = report.add_paragraph(text = f'{str(student.name).upper()} {str(student.surname).upper()}')
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font = p4.runs[0].font
        font.size = Pt(18)
        report.add_paragraph('')
        p5 = report.add_paragraph(text = 'Oceny:  ')
        font = p5.runs[0].font
        font.color.rgb = RGBColor(0, 191, 255)  # RGB color for light blue
        grades = SingleGrade.objects.filter(student=student).order_by('date')
        for grade in grades:
            r = p5.add_run(str(grade.grade) + ', ')
            r.font.color.rgb = RGBColor(0, 191, 255)
        report.add_paragraph('')
        p6 = report.add_paragraph(text = 'Daty nieobecności:  ')
        font = p6.runs[0].font
        font.color.rgb = RGBColor(250, 128, 114)  # RGB color for salmon
        lessons_absent = student.lessons_absent.all()
        for lesson in lessons_absent:
            r = p6.add_run(lesson.date.strftime('%d-%m') + ', ')
            r.font.color.rgb = RGBColor(250, 128, 114)
        report.add_paragraph('')
        notes = PersonalNote.objects.filter(student=student).order_by('date')
        p7 = report.add_paragraph(text = 'Uwagi lektora:')
        font = p7.runs[0].font
        font.color.rgb = RGBColor(0, 128, 0)  # RGB color for a non-vibrant green
        for note in notes:
            p7 = report.add_paragraph(text = note.note)
            font = p7.runs[0].font
            font.color.rgb = RGBColor(0, 128, 0)  # RGB color for a non-vibrant green
        
        report.add_page_break()
	
		
    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = f'inline; filename="quaterly report for group {group}.docx"'
    report.save(response)
    return response 

